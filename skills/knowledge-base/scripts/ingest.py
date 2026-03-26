#!/usr/bin/env python3
"""
文档摄入模块 - 完整流程：解析→分类→切片→索引
"""

import os
import sys
import json
import argparse
import hashlib
import shutil
import re
from pathlib import Path
from datetime import datetime

# 添加脚本目录到路径
script_dir = os.path.dirname(__file__)
sys.path.insert(0, script_dir)

# 首先检查并安装依赖（包含虚拟环境自动创建）
try:
    from dependency_manager import ensure_deps, is_using_venv, get_venv_info
    deps_ok = ensure_deps()
    if not deps_ok:
        print("❌ 依赖安装失败，无法继续归档")
        sys.exit(1)
    
    # 如果依赖刚安装好，但当前不在虚拟环境中，提示用户重新运行
    if not is_using_venv():
        _, venv_python = get_venv_info()
        script_path = os.path.abspath(__file__)
        print(f"\n⚠️  当前未使用虚拟环境，但依赖已就绪")
        print(f"   请使用虚拟环境重新运行:")
        print(f"   {venv_python} {script_path} {' '.join(sys.argv[1:])}")
        sys.exit(1)
except Exception as e:
    print(f"⚠️  依赖检查失败: {e}")
    print("   继续尝试执行...")

import numpy as np
from embedder import encode_chunks
from classifier import AutoClassifier
from notebook_manager import NotebookManager

class DocumentIngester:
    """文档摄入器"""
    
    def __init__(self, kb_base_path: str = None):
        if kb_base_path is None:
            kb_base_path = os.path.expanduser("~/.openclaw/workspace/docs-db")
        self.kb_base_path = kb_base_path
        self.classifier = AutoClassifier(kb_base_path)
        self.manager = NotebookManager(kb_base_path)
    
    def _calculate_file_hash(self, file_path: str) -> str:
        """计算文件MD5哈希"""
        hash_md5 = hashlib.md5()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_md5.update(chunk)
        return hash_md5.hexdigest()
    
    def _check_duplicate(self, file_hash: str, file_name: str) -> dict:
        """
        检查文档是否已存在
        
        Returns:
            如果重复返回 {"is_duplicate": True, "notebook": "项目名", "doc_id": "xxx"}
            否则返回 {"is_duplicate": False}
        """
        # 遍历所有notebook
        for notebook_name in self.manager.list_notebooks():
            chunks_file = self.manager.get_chunks_path(notebook_name)
            if not os.path.exists(chunks_file):
                continue
            
            # 读取chunks.jsonl，检查是否有相同文件名的文档
            try:
                with open(chunks_file, 'r', encoding='utf-8') as f:
                    for line in f:
                        chunk = json.loads(line.strip())
                        existing_source = chunk.get("source", "")
                        existing_doc_id = chunk.get("doc_id", "")
                        
                        # 检查文件名是否相同
                        if existing_source == file_name:
                            return {
                                "is_duplicate": True,
                                "notebook": notebook_name,
                                "doc_id": existing_doc_id,
                                "reason": "filename"
                            }
            except Exception:
                continue
        
        return {"is_duplicate": False}
    
    def parse_document(self, file_path: str) -> tuple:
        """解析文档，返回(内容, 标题, markdown内容)"""
        file_path = os.path.expanduser(file_path)
        title = os.path.basename(file_path)
        ext = Path(file_path).suffix.lower()
        
        if ext == '.pdf':
            content, markdown = self._parse_pdf(file_path)
        elif ext in ['.docx', '.doc']:
            content, markdown = self._parse_docx(file_path)
        elif ext in ['.txt', '.md', '.markdown']:
            content = self._parse_text(file_path)
            markdown = content  # 文本文件直接作为markdown
        else:
            # 尝试作为文本读取
            try:
                content = self._parse_text(file_path)
                markdown = content
            except:
                raise ValueError(f"不支持的文件格式: {ext}")
        
        return content, title, markdown
    
    def _parse_pdf(self, file_path: str) -> tuple:
        """解析PDF，返回(纯文本内容, markdown格式内容)"""
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(file_path)
            
            text_parts = []
            markdown_parts = []
            
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    text_parts.append(text)
                    # 生成markdown格式，保留页码信息
                    markdown_parts.append(f"## 第 {i+1} 页\n\n{text}")
            
            plain_text = "\n\n".join(text_parts)
            markdown_text = f"# {os.path.basename(file_path)}\n\n" + "\n\n".join(markdown_parts)
            
            return plain_text, markdown_text
        except ImportError:
            print("⚠️  PyPDF2未安装，尝试纯文本读取")
            text = self._parse_text(file_path)
            return text, text
    
    def _parse_docx(self, file_path: str) -> tuple:
        """解析Word，返回(纯文本内容, markdown格式内容)"""
        try:
            from docx import Document
            from docx.enum.text import WD_STYLE_TYPE
            
            doc = Document(file_path)
            
            markdown_parts = [f"# {os.path.basename(file_path)}\n"]
            text_parts = []
            
            for para in doc.paragraphs:
                if not para.text.strip():
                    continue
                
                text_parts.append(para.text)
                
                # 根据样式判断标题级别
                style_name = para.style.name if para.style else "Normal"
                
                if style_name.startswith('Heading 1') or style_name.startswith('标题 1'):
                    markdown_parts.append(f"\n## {para.text}\n")
                elif style_name.startswith('Heading 2') or style_name.startswith('标题 2'):
                    markdown_parts.append(f"\n### {para.text}\n")
                elif style_name.startswith('Heading 3') or style_name.startswith('标题 3'):
                    markdown_parts.append(f"\n#### {para.text}\n")
                else:
                    # 普通段落
                    # 检测列表
                    text = para.text.strip()
                    if text.startswith(('•', '·', '-', '–', '*')):
                        markdown_parts.append(f"- {text[1:].strip()}")
                    elif re.match(r'^\d+[.、]\s', text):
                        markdown_parts.append(f"1. {re.sub(r'^\d+[.、]\\s*', '', text)}")
                    else:
                        markdown_parts.append(text)
            
            # 处理表格
            for table in doc.tables:
                markdown_parts.append("\n")
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    markdown_parts.append(f"| {row_text} |")
                markdown_parts.append("")
            
            plain_text = "\n\n".join(text_parts)
            markdown_text = "\n".join(markdown_parts)
            
            return plain_text, markdown_text
        except ImportError:
            print("⚠️  python-docx未安装，尝试纯文本读取")
            text = self._parse_text(file_path)
            return text, text
    
    def _parse_text(self, file_path: str) -> str:
        """解析文本文件"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    
    def chunk_text(self, text: str, chunk_size: int = 256, overlap: int = 50) -> list:
        """将文本切分为重叠的chunk"""
        chunks = []
        start = 0
        text_len = len(text)
        
        while start < text_len:
            end = start + chunk_size
            chunk = text[start:end]
            
            # 尽量在句子边界切分
            if end < text_len:
                # 向后找句号、换行或空格
                for i in range(min(50, len(chunk)), 0, -1):
                    if chunk[-i] in '。！？\n ':
                        chunk = chunk[:-i]
                        break
            
            if chunk.strip():
                chunks.append(chunk.strip())
            
            start += chunk_size - overlap
        
        return chunks
    
    def ingest(self, file_path: str, notebook: str = None, auto_classify: bool = True) -> dict:
        """
        摄入文档
        
        Args:
            file_path: 文档路径
            notebook: 指定notebook（None则自动分类）
            auto_classify: 是否自动分类
        
        Returns:
            摄入结果信息
        """
        file_path = os.path.expanduser(file_path)
        file_name = os.path.basename(file_path)
        
        print(f"📥 正在处理: {file_path}")
        
        # 1. 检查重复文档
        duplicate_check = self._check_duplicate("", file_name)
        if duplicate_check["is_duplicate"]:
            print(f"⚠️  文档已存在知识库中！")
            print(f"   📁 所在项目: {duplicate_check['notebook']}")
            print(f"   📝 文档名称: {file_name}")
            print(f"   💡 无需重复归档")
            return {
                "status": "skipped",
                "reason": "duplicate",
                "notebook": duplicate_check["notebook"],
                "message": f"文档 '{file_name}' 已存在于项目 '{duplicate_check['notebook']}' 中，无需重复归档"
            }
        
        # 2. 解析文档
        content, title, markdown = self.parse_document(file_path)
        print(f"📄 解析完成: {len(content)} 字符")
        
        # 2. 确定目标notebook
        if notebook:
            # 用户指定
            target_notebook = notebook
            decision = "user_specified"
            similarity = 1.0
            
            # 如果不存在则创建
            if not self.manager.notebook_exists(target_notebook):
                self.manager.create_notebook(target_notebook)
                print(f"✨ 创建新Notebook: {target_notebook}")
        elif auto_classify:
            # 自动分类
            profile = self.classifier.extract_document_profile(content, title)
            target_notebook, decision, similarity = self.classifier.classify(profile)
            
            if decision == "create_new":
                self.manager.create_notebook(target_notebook)
                print(f"✨ 创建新Notebook: {target_notebook}")
            
            # 打印分类详情
            all_sims = self.classifier.get_all_similarities(profile)
            if all_sims:
                print("🔍 相似度排名:")
                for nb, sim in all_sims[:3]:
                    print(f"   {'✓' if nb == target_notebook else ' '} {nb}: {sim:.1%}")
        else:
            target_notebook = "default"
            decision = "default"
            similarity = 0.0
            if not self.manager.notebook_exists(target_notebook):
                self.manager.create_notebook(target_notebook)
        
        # 3. 创建新的目录结构
        doc_id = hashlib.md5(f"{file_path}{datetime.now()}".encode()).hexdigest()[:8]
        ext = Path(file_path).suffix
        
        # 获取基础文档目录
        base_docs_path = self.manager.get_documents_path(target_notebook)
        
        # 创建 raw 和 text 子目录
        raw_dir = os.path.join(base_docs_path, "raw")
        text_dir = os.path.join(base_docs_path, "text")
        os.makedirs(raw_dir, exist_ok=True)
        os.makedirs(text_dir, exist_ok=True)
        
        # 保存原始文档到 raw/
        raw_filename = f"{doc_id}_{title}"
        raw_dest_path = os.path.join(raw_dir, raw_filename)
        shutil.copy2(file_path, raw_dest_path)
        
        # 保存markdown版本到 text/
        md_filename = f"{doc_id}_{Path(title).stem}.md"
        md_dest_path = os.path.join(text_dir, md_filename)
        with open(md_dest_path, 'w', encoding='utf-8') as f:
            f.write(markdown)
        
        print(f"📁 原始文档保存至: documents/raw/{raw_filename}")
        print(f"📝 Markdown保存至: documents/text/{md_filename}")
        
        # 4. 切分chunk（使用纯文本内容）
        chunks = self.chunk_text(content)
        print(f"✂️  切分为 {len(chunks)} 个chunk")
        
        # 5. 生成向量
        print("🔄 正在生成向量索引...")
        embeddings = encode_chunks(chunks, show_progress=True)
        
        # 6. 生成TF-IDF
        from sklearn.feature_extraction.text import TfidfVectorizer
        tfidf = TfidfVectorizer(max_features=5000)
        tfidf_matrix = tfidf.fit_transform(chunks)
        
        # 7. 保存索引
        index_path = self.manager.get_index_path(target_notebook)
        os.makedirs(index_path, exist_ok=True)
        
        # 保存chunks元数据
        chunks_data = []
        for i, chunk in enumerate(chunks):
            chunks_data.append({
                "id": f"{doc_id}_{i}",
                "text": chunk,
                "source": title,
                "doc_id": doc_id,
                "chunk_index": i,
                "notebook": target_notebook,
                "raw_path": f"documents/raw/{raw_filename}",
                "text_path": f"documents/text/{md_filename}"
            })
        
        # 追加或创建chunks.jsonl
        chunks_file = os.path.join(self.manager.get_notebook_path(target_notebook), "chunks.jsonl")
        with open(chunks_file, 'a', encoding='utf-8') as f:
            for item in chunks_data:
                f.write(json.dumps(item, ensure_ascii=False) + '\n')
        
        # 保存向量（追加模式）
        embeddings_file = os.path.join(index_path, "embeddings.npy")
        if os.path.exists(embeddings_file):
            existing = np.load(embeddings_file)
            all_embeddings = np.vstack([existing, embeddings])
        else:
            all_embeddings = embeddings
        np.save(embeddings_file, all_embeddings)
        
        # 保存TF-IDF（简化：每次重建）
        import pickle
        with open(os.path.join(index_path, "tfidf_vectorizer.pkl"), 'wb') as f:
            pickle.dump(tfidf, f)
        
        # 更新TF-IDF矩阵（追加）
        tfidf_file = os.path.join(index_path, "tfidf_matrix.npz")
        from scipy import sparse
        if os.path.exists(tfidf_file):
            existing_tfidf = sparse.load_npz(tfidf_file)
            all_tfidf = sparse.vstack([existing_tfidf, tfidf_matrix])
        else:
            all_tfidf = tfidf_matrix
        sparse.save_npz(tfidf_file, all_tfidf)
        
        print(f"✅ 归档完成!")
        print(f"   📁 Notebook: {target_notebook}")
        print(f"   📝 文档: {title}")
        print(f"   🧩 Chunks: {len(chunks)}")
        
        return {
            "doc_id": doc_id,
            "notebook": target_notebook,
            "decision": decision,
            "similarity": similarity,
            "chunks": len(chunks),
            "raw_path": raw_dest_path,
            "text_path": md_dest_path
        }

# CLI入口
if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="文档摄入工具")
    parser.add_argument("file", help="文档路径")
    parser.add_argument("--notebook", "-n", help="指定归档的notebook")
    parser.add_argument("--auto-classify", "-a", action="store_true", default=True, help="自动分类")
    
    args = parser.parse_args()
    
    ingester = DocumentIngester()
    result = ingester.ingest(args.file, args.notebook, args.auto_classify)
    
    print("\n📊 摄入结果:")
    print(json.dumps(result, ensure_ascii=False, indent=2))
